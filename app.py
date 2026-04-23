"""
재정경제부장관  국회 질의 공식 답변 초안 작성 (Streamlit)
근거: `국회/code` 폴더의 프리셋 질의서 + `국회/참고자료` 폴더 문서.

실행 (저장소 루트에서):
  uv run streamlit run 국회/code/app.py

환경: OPENAI_API_KEY (로컬은 AI-Education/.env, Streamlit Cloud는 Secrets)
"""

from __future__ import annotations

import hashlib
import os
import re
from pathlib import Path
from typing import Any

import streamlit as st
from dotenv import load_dotenv
from langchain_community.document_loaders import PyPDFLoader
from langchain_community.vectorstores import FAISS
from langchain_core.documents import Document
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter

# ---------------------------------------------------------------------------
APP_DIR = Path(__file__).resolve().parent
REPO_ROOT = APP_DIR.parent.parent
ASSEMBLY_ROOT = REPO_ROOT / "국회"
REF_DIR = ASSEMBLY_ROOT / "참고자료"


def _first_existing(paths: list[Path]) -> Path | None:
    for p in paths:
        if p.is_file():
            return p
    return None


LOGO_PATH = _first_existing([APP_DIR / "logo.png", REPO_ROOT / "logo.png"]) or (APP_DIR / "logo.png")
LOGO_HEADER_WIDTH_PX = 480

PRESET_QUERY_DB: list[Path] = [
    APP_DIR / "260402 재경위 전체회의 대비.txt",
    APP_DIR / "260402 재경위 현안질의.txt",
    APP_DIR / "260407 예결위 종합정책질의(1+2일차).txt",
]

CHUNK_SIZE = 900
CHUNK_OVERLAP = 150
RETRIEVE_K = 8
MMR_FETCH_K = 24
MMR_LAMBDA = 0.55
MODEL = "gpt-4o-mini"
EMBED_MODEL = "text-embedding-3-small"

SUPPORTED_REF_EXT = {".txt", ".md", ".pdf", ".docx"}


def load_env() -> None:
    """로컬: `.env` 로드. Streamlit Cloud: `st.secrets`의 값을 환경변수로 승격."""
    for env_path in (REPO_ROOT / ".env", APP_DIR / ".env"):
        if env_path.is_file():
            load_dotenv(env_path)
            break
    try:
        for key in ("OPENAI_API_KEY",):
            if not os.getenv(key) and key in st.secrets:
                os.environ[key] = str(st.secrets[key])
    except Exception:
        pass


def has_openai() -> bool:
    load_env()
    return bool(os.getenv("OPENAI_API_KEY", "").strip())


def read_txt_like(path: Path) -> str:
    for enc in ("utf-8-sig", "utf-8", "cp949"):
        try:
            return path.read_text(encoding=enc)
        except UnicodeDecodeError:
            continue
    return path.read_text(encoding="utf-8", errors="replace")


def read_pdf_text(path: Path) -> str:
    loader = PyPDFLoader(str(path))
    docs = loader.load()
    return "\n\n".join(d.page_content for d in docs if d.page_content)


def read_docx_text(path: Path) -> str:
    from docx import Document as DocxDocument

    doc = DocxDocument(str(path))
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def load_document_text(path: Path) -> str:
    ext = path.suffix.lower()
    if ext in (".txt", ".md"):
        return read_txt_like(path)
    if ext == ".pdf":
        return read_pdf_text(path)
    if ext == ".docx":
        return read_docx_text(path)
    raise ValueError(f"지원하지 않는 형식입니다: {path}")


def list_files_by_mtime(directory: Path, extensions: set[str]) -> list[Path]:
    if not directory.is_dir():
        return []
    out: list[Path] = []
    for p in directory.rglob("*"):
        if p.is_file() and p.suffix.lower() in extensions:
            out.append(p)
    out.sort(key=lambda x: x.stat().st_mtime, reverse=True)
    return out


def load_preset_query_documents() -> list[Document]:
    """`국회/code`에 넣어둔 프리셋 질의서(3개)를 RAG 검색용 Document로 로드."""
    docs: list[Document] = []
    for p in PRESET_QUERY_DB:
        if not p.is_file():
            continue
        try:
            text = load_document_text(p)
        except Exception as e:
            docs.append(
                Document(
                    page_content=f"[질의서 로드 실패: {p.name} — {e}]",
                    metadata={"source": f"질의서:{p.name}", "filename": p.name, "error": True},
                )
            )
            continue
        if not text.strip():
            continue
        docs.append(
            Document(
                page_content=text.strip(),
                metadata={"source": f"질의서:{p.name}", "filename": p.name, "kind": "질의서"},
            )
        )
    return docs


def load_reference_documents() -> list[Document]:
    paths = list_files_by_mtime(REF_DIR, SUPPORTED_REF_EXT)
    docs: list[Document] = []
    for p in paths:
        try:
            text = load_document_text(p)
        except Exception as e:
            docs.append(
                Document(
                    page_content=f"[파일 로드 실패: {p.name} — {e}]",
                    metadata={"source": str(p), "error": True},
                )
            )
            continue
        if not text.strip():
            continue
        docs.append(
            Document(
                page_content=text.strip(),
                metadata={"source": str(p), "filename": p.name},
            )
        )
    return docs


@st.cache_resource
def build_faiss_index(_docs_fingerprint: str, docs_data: tuple[tuple[str, dict[str, Any]], ...]) -> FAISS | None:
    if not has_openai() or not docs_data:
        return None
    docs = [Document(page_content=t, metadata=m) for t, m in docs_data]
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        length_function=len,
    )
    splits = splitter.split_documents(docs)
    if not splits:
        return None
    emb = OpenAIEmbeddings(model=EMBED_MODEL)
    return FAISS.from_documents(splits, emb)


def fingerprint_docs(docs: list[Document]) -> str:
    """프로세스 간 안정적인 캐시 키용(내용 해시는 hash() 대신 SHA-256)."""
    parts = []
    for d in docs:
        src = d.metadata.get("source", "")
        h = hashlib.sha256(d.page_content.encode("utf-8", errors="replace")).hexdigest()[:16]
        parts.append(f"{src}:{h}")
    return "|".join(parts)


def retrieve_context(store: FAISS | None, query: str) -> str:
    if store is None or not query.strip():
        return ""
    q = query.strip()
    try:
        fetch_k = min(MMR_FETCH_K, max(RETRIEVE_K * 2, 10))
        docs = store.max_marginal_relevance_search(
            q,
            k=RETRIEVE_K,
            fetch_k=fetch_k,
            lambda_mult=MMR_LAMBDA,
        )
    except Exception:
        docs = [d for d, _ in store.similarity_search_with_score(q, k=RETRIEVE_K)]
    blocks: list[str] = []
    for doc in docs:
        src = doc.metadata.get("filename") or doc.metadata.get("source", "")
        blocks.append(f"— 출처: {src}\n{doc.page_content.strip()}")
    return "\n\n".join(blocks)


SUMMARY_SYSTEM = """당신은 대한민국 기획재정부 국회 대관 업무를 보좌하는 전문가입니다.
주어진 질의서 전문과 질의 의원 이름만을 바탕으로, 의원이 제기한 핵심 문제점과 요구사항을 객관적으로 요약합니다.
추측으로 새 사실을 만들지 마십시오. 질의서에 없는 내용은 '질의서에 명시되지 않음'으로 표시합니다.
출력은 한국어이며, 불릿 포인트 위주로 간결하게 작성합니다."""

DRAFT_SYSTEM = """당신은 대한민국 기획재정부(재정경제부) 장관  국회 서면·구두 답변 초안을 작성하는 필기관입니다.

[톤과 문체]
- 격식 있는 공문서체(하십시오체)를 사용합니다.
- 장관의 품격에 맞는 담담하고 책임 있는 표현을 씁니다.

[질의 반영 — 반복 금지]
- 이번에 주어진 '질의서 전문'의 쟁점·용어·항목(가·나·다, 번호 목록 등)을 반드시 짚고, 일반론·상투구만으로 뭉뚱그려 답하지 마십시오.
- 질의서마다 [세부 답변]의 소제목과 논거 배치가 달라야 합니다. 이전 답변과 동일한 문단을 재사용하지 마십시오.

[사실과 근거]
- 아래에 제공된 '검색 근거 발췌'에 명시된 내용만 수치·정책 사실로 단정합니다.
- 근거에 없는 구체적 수치·연도·사업명은 만들지 말고, '관계 부처와 협의하여 검토 중이오니 향후 별도로 보고드리겠습니다' 또는 '공식 통계·결산 공개 자료를 근거로 향후 보고드리겠습니다' 등으로 처리합니다.
- 질의서와 근거가 충돌하면 보수적으로 서술하고, 추가 확인이 필요함을 밝힙니다.

[구성]
반드시 다음 세 부제를 순서대로 사용합니다(각 부제는 한 줄에 굵게 쓸 수 있도록 대괄호 포함).

[답변 요지]
[세부 답변]
[향후 계획]

- 절대로 [인사] 부제나 '존경하는 ○○ 의원님' 같은 인사말·감사말 문단을 넣지 마십시오.
- 본문은 [답변 요지]에서 두괄식으로 곧바로 시작합니다.
- [세부 답변]에서는 질의 사항을 나누어 소제목을 달고, 가능하면 불릿 포인트로 가독성을 높입니다.

[금지]
- 문서 근거 없는 자의적 통계 창작
- 정치적 공격·비방 표현
- 인사말·사의 표명 문구(의정활동에 경의, 질의에 감사, 귀하 등)

[특수 규칙 — 부총리(장관)의 선거 출마 관련 질의]
- 질의 내용이 부총리(재정경제부장관·기획재정부장관)의 향후 선거(대선·총선·지방선거 등) **출마 여부·정치적 행보·정당 활동·대권 관련 의사**를 묻는 경우에는, [세부 답변]에서 정책적 논거를 장황히 나열하지 말고 다음 취지를 간명한 공문서체로 답하십시오.
  - 현재 부총리는 오로지 **경제와 민생 안정**에 전념하고 있음.
  - **선거 출마·정치적 거취는 전혀 고려하고 있지 않음**을 분명히 함.
  - 따라서 해당 사안은 답변의 대상이 되지 않으며, 맡은 바 **국정과 재정운용에 책임을 다할 것**임을 밝힘.
- 이 규칙이 적용될 때는 [답변 요지]도 위 취지(경제·민생 전념, 선거 고려하지 않음)를 두괄식으로 제시하고, [향후 계획]은 경제정책 관리에 대한 책임 이행으로 간결히 마무리합니다.
- 단, 질의가 부총리의 선거 출마가 아닌 일반 선거제도·정치자금·공직선거법 등 정책·제도적 쟁점을 다루는 경우에는 이 규칙을 적용하지 않고, 통상의 정책 답변 절차를 따릅니다."""


def run_summary(llm: ChatOpenAI, mp_name: str, query_full: str) -> str:
    msg = (
        f"질의 의원: {mp_name}\n\n"
        f"=== 질의서 전문 ===\n{query_full.strip()}\n=== 끝 ===\n\n"
        "다음을 작성하십시오.\n"
        "1) 핵심 문제의식(의원이 지적하는 쟁점)\n"
        "2) 구체적 요구사항(항목별)\n"
        "3) 답변 시 유의할 점(질의서에 명시된 조건·근거 요청 등)"
    )
    r = llm.invoke(
        [
            ("system", SUMMARY_SYSTEM),
            ("human", msg),
        ]
    )
    return (r.content or "").strip()


def query_fingerprint(text: str) -> str:
    return hashlib.sha256(text.strip().encode("utf-8", errors="replace")).hexdigest()[:12]


def _banner_title_html() -> str:
    """브라우저가 어중간한 음절에서 끊지 않도록 스타일 + 의도된 줄바꿈(대시 뒤)."""
    line1 = "국회 질의 대응 —"
    line2 = "재정경제부장관 답변 초안"
    return (
        f'<h1 class="mof-banner-title">{line1}<br/>{line2}</h1>'
    )


def render_title_banner() -> None:
    """상단: 제목(좌) + 저장소 루트 `logo.png`(우측)."""
    cap = (
        "사용자의 질문에 대해 장관 답변 초안을 생성합니다. 제출 전 반드시 담당자 검토가 필요합니다."
    )
    st.markdown(
        """
        <style>
        .mof-banner-title {
            font-size: clamp(1.35rem, 2.4vw, 1.95rem);
            font-weight: 700;
            line-height: 1.38;
            letter-spacing: -0.02em;
            margin: 0 0 0.55rem 0;
            padding: 0;
            word-break: keep-all;
            line-break: strict;
            overflow-wrap: break-word;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )

    if LOGO_PATH.is_file():
        try:
            left, right = st.columns([3, 2], vertical_alignment="center")
        except TypeError:
            left, right = st.columns([3, 2])
        with left:
            st.markdown(_banner_title_html(), unsafe_allow_html=True)
            st.caption(cap)
        with right:
            st.image(str(LOGO_PATH), width=LOGO_HEADER_WIDTH_PX)
    else:
        st.markdown(_banner_title_html(), unsafe_allow_html=True)
        st.caption(cap)


def run_draft(
    llm: ChatOpenAI,
    mp_name: str,
    query_full: str,
    summary: str,
    context: str,
) -> str:
    ctx = context.strip() or "(참고자료에서 검색된 근거가 없습니다. 일반적이고 보수적인 표현만 사용하십시오.)"
    qfp = query_fingerprint(query_full)
    human = (
        f"질의 의원: {mp_name}\n"
        f"답변 주체: 재정경제부장관(기획재정부장관)\n"
        f"질의 지문 식별자(참고): {qfp}\n\n"
        f"=== 현안 요약(내부 참고) ===\n{summary}\n\n"
        f"=== 질의서 전문 ===\n{query_full.strip()}\n=== 끝 ===\n\n"
        f"=== 검색 근거 발췌 (이 외의 수치·사실을 새로 만들지 마십시오) ===\n{ctx}\n"
        "=== 끝 ===\n\n"
        "위 질의서 전문과 식별자에 맞추어 장관  답변 초안 전문을 작성하십시오. "
        "질의 지문이 바뀌면 내용도 그에 맞게 달라져야 합니다."
    )
    r = llm.invoke(
        [
            ("system", DRAFT_SYSTEM),
            ("human", human),
        ]
    )
    return (r.content or "").strip()


def main() -> None:
    _icon = str(LOGO_PATH) if LOGO_PATH.is_file() else "📋"
    st.set_page_config(
        page_title="국회 질의 — 장관 답변 초안",
        page_icon=_icon,
        layout="wide",
    )
    render_title_banner()

    with st.sidebar:
        st.header("사용 방법 안내")
        st.markdown(
            "1. **질의 의원 성명**을 입력합니다.\n"
            "2. **질문 내용**에 답변이 필요한 질의를 자유롭게 적습니다.\n"
            "3. **「답변 초안 생성」** 버튼을 누릅니다.\n"
            "4. **질의 DB**에서 관련 내용을 찾아, \n"
            "   상단에 **장관 답변 초안**, 하단에 **현안 요약**이 표시됩니다.\n"
            "5. 필요 시 **초안 텍스트 다운로드**로 저장합니다.\n\n"
            "※ 답변은 AI 초안이며, 제출 전 반드시 담당자의 **사실 확인·문장 검토**가 필요합니다."
        )

    load_env()
    if not has_openai():
        st.error(
            "`OPENAI_API_KEY`가 없습니다. 로컬은 `.env`, Streamlit Cloud는 Secrets에 설정하십시오."
        )
        return

    preset_docs = load_preset_query_documents()
    ref_docs = load_reference_documents()
    all_docs = preset_docs + ref_docs

    if not preset_docs:
        st.warning(
            "프리셋 질의서 파일을 찾지 못했습니다. `국회/code` 폴더에 다음 파일을 두십시오:\n- "
            + "\n- ".join(p.name for p in PRESET_QUERY_DB)
        )

    fp = fingerprint_docs(all_docs)
    docs_tuple = tuple((d.page_content, dict(d.metadata)) for d in all_docs)
    store = build_faiss_index(fp, docs_tuple)
    if all_docs and store is None:
        st.warning("질의 DB가 있으나 벡터 인덱스를 만들지 못했습니다. API 키를 확인하십시오.")

    mp_name = st.text_input("질의 의원 성명", placeholder="예: 홍길동")
    user_question = st.text_area(
        "질문 내용",
        height=160,
        placeholder="예) 최근 물가 상승 대응을 위한 정부의 긴급 조치와 재원 확보 방안을 밝혀 주시기 바랍니다.",
        help="질의 DB에서 이 질문과 관련된 내용을 검색해 답변 초안을 만듭니다.",
    )

    go = st.button("답변 초안 생성", type="primary", use_container_width=True)
    if not go:
        return

    if not (mp_name or "").strip():
        st.error("질의 의원 성명을 입력하십시오.")
        return
    if not (user_question or "").strip():
        st.error("질문 내용을 입력하십시오.")
        return
    if store is None:
        st.error("질의 DB가 비어 있어 답변을 생성할 수 없습니다.")
        return

    query_text = user_question.strip()

    llm_summary = ChatOpenAI(model=MODEL, temperature=0.2)
    llm_draft = ChatOpenAI(model=MODEL, temperature=0.45)

    with st.spinner("1/2 현안 파악(질문 요약, 내부 참고)…"):
        summary = run_summary(llm_summary, mp_name.strip(), query_text)

    retrieval_query = f"{mp_name} {query_text}"
    with st.spinner("2/2 질의 DB 검색 및 답변 초안 작성…"):
        ctx = retrieve_context(store, retrieval_query)
        draft = run_draft(llm_draft, mp_name.strip(), query_text, summary, ctx)

    st.subheader("1. 재정경제부장관 답변 초안")
    st.markdown(draft)

    st.subheader("2. 현안 파악 (질문 기준)")
    st.markdown(summary)

    with st.expander("질의 DB에서 검색된 근거 발췌 (일부)", expanded=False):
        st.markdown(ctx or "_검색 결과 없음_")

    st.download_button(
        "초안 텍스트 다운로드 (.txt)",
        data=f"[답변 초안]\n\n{draft}\n\n---\n\n[현안 요약]\n\n{summary}",
        file_name=f"장관답변초안_{re.sub(r'[^가-힣a-zA-Z0-9_]+', '_', mp_name.strip())}.txt",
        mime="text/plain; charset=utf-8",
    )


if __name__ == "__main__":
    main()
